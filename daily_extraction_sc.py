from selenium import webdriver
from selenium.webdriver.edge.service import Service as EdgeService
from webdriver_manager.microsoft import EdgeChromiumDriverManager
from selenium.webdriver.edge.options import Options as EdgeOptions
from selenium.webdriver.common.by import By
from bs4 import BeautifulSoup
from docx import Document

import datetime
import time
import pandas as pd
import requests

# 设置 EdgeOptions
options = EdgeOptions()
options.add_argument("--headless")  # 无头模式
# 初始化 WebDriver
driver = webdriver.Edge(service=EdgeService(EdgeChromiumDriverManager().install()), options=options)

today = datetime.datetime.now().strftime("%m-%d")
yesterday = (datetime.datetime.now() - datetime.timedelta(days=1)).strftime("%m-%d")
companies = pd.read_csv("data/sichuan_new.csv")

current_month_day = datetime.datetime.now().strftime("%m-%d")

# 创建一个 Word 文档
doc = Document()
doc.add_heading(f'四川上市公司资讯 - {current_month_day}', level=1)

for index, row in companies.iterrows():
    # 打开 dataweb 的网址
    print(f"Source: {row["website"]}")
    driver.get(row["website"])
    time.sleep(3)  # 等待页面加载

    # 使用 BeautifulSoup 解析页面内容
    soup = BeautifulSoup(driver.page_source, "html.parser")
    qnb_list = soup.find('div', class_="qnb_list")
    for li in qnb_list.find_all("li"):
        time_span = li.find('span', class_='time')
        if time_span:
            time_text = time_span.text.strip()
            month_day = datetime.datetime.strptime(time_text, '%m-%d').strftime('%m-%d')
        if month_day == current_month_day:
            a_tag = li.find('a')
            if a_tag:
                href = a_tag.get('href')
                content = a_tag.text.strip()
                print(f'Href: {href}, Content: {content}')

                try:
                    response = requests.get(href)
                    response.raise_for_status()  # 检查请求是否成功
                    href_soup = BeautifulSoup(response.text, 'html.parser')

                    # 提取 <title>
                    title = href_soup.find('title')
                    title_text = title.text.strip() if title else "No Title"
                    # print(f'Title: {title_text}')

                    # 提取 <meta name="keywords">
                    meta_keywords = href_soup.find('meta', attrs={'name': 'keywords'})
                    keywords_content = meta_keywords.get('content', '').strip() if meta_keywords else "No Keywords"
                    # print(f'Keywords: {keywords_content}')

                    # 提取 <meta name="description">
                    meta_description = href_soup.find('meta', attrs={'name': 'description'})
                    description_content = meta_description.get('content', '').strip() if meta_description else "No Description"
                    # print(f'Description: {description_content}')

                    # 将信息写入 Word 文档
                    doc.add_paragraph(f'链接: {href}')  # 链接
                    doc.add_paragraph(f'标题: {title_text}')  # 标题
                    doc.add_paragraph(f'关键词: {keywords_content}')  # 关键词
                    doc.add_paragraph(f'描述: {description_content}')  # 描述
                    doc.add_paragraph('————' * 8)  # 中文分隔线

                except requests.RequestException as e:
                    print(f'Failed to retrieve {href}: {e}')

# 保存 Word 文档
doc.save(f'daily/四川上市公司资讯_{current_month_day}.docx')